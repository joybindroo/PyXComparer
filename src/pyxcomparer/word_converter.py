"""XLSForm to Word document converter module.

This module provides functionality to convert ODK XLSForm metadata (YAML) 
into a human-readable Word document for review and documentation.
"""

from pathlib import Path
from typing import Optional
import yaml
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_yaml_to_word(
    yaml_path: Path | str, 
    output_path: Optional[Path | str] = None
) -> Path:
    """
    Convert YAML representation of an XLSForm to a detailed Word document.

    Args:
        yaml_path: Path to the input YAML file.
        output_path: Path for the output .docx file. If None, generated based on yaml_path.

    Returns:
        Path to the generated Word document.
    """
    yaml_path = Path(yaml_path)
    if output_path is None:
        output_path = yaml_path.with_suffix(".docx")
    else:
        output_path = Path(output_path)

    with open(yaml_path, 'r', encoding='utf-8') as f:
        # Handle YAML files written as a series of single-item lists
        raw_content = f.read()
        data = []
        for block in raw_content.split('\n---\n'):
            try:
                parsed = yaml.safe_load(block)
                if isinstance(parsed, list):
                    data.extend(parsed)
                elif isinstance(parsed, dict):
                    data.append(parsed)
            except yaml.YAMLError:
                continue
        
        if not data:
            with open(yaml_path, 'r', encoding='utf-8') as f2:
                data = list(yaml.safe_load_all(f2))
                flattened = []
                for item in data:
                    if isinstance(item, list):
                        flattened.extend(item)
                    elif isinstance(item, dict):
                        flattened.append(item)
                data = flattened

    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = doc.add_heading('ODK XLSForm - Technical Specification', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    form_title = "XLSForm"
    for item in data:
        if isinstance(item, dict) and 'form_title' in item:
            form_title = item['form_title']
            break
    
    doc.add_paragraph(f"Form Title: {form_title}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    choice_map = {}
    for item in data:
        if not isinstance(item, dict): continue
        for key, value in item.items():
            if key.startswith('Option_List::'):
                list_name = key.split('::')[1]
                choice_map[list_name] = value

    current_group = None

    for item in data:
        if not isinstance(item, dict): continue
        
        q_type = item.get('type', '').lower()
        name = item.get('name', '')
        label_en = item.get('label::English', '')
        label_hi = item.get('label::Hindi', '')
        
        if q_type == 'begin group':
            current_group = label_en if label_en else name
            doc.add_heading(f"Section: {current_group}", level=1)
            continue
        elif q_type == 'end group':
            current_group = None
            doc.add_paragraph("-" * 30)
            continue
            
        if q_type in ['audit', 'start', 'end', 'deviceid', 'phonenumber', 'today', 'background-audio']:
            continue
            
        if not label_en and not label_hi and not name:
            continue

        full_label = f"{label_en}\n{label_hi}" if label_hi else label_en
        if not full_label:
            full_label = f"Question: {name}"

        p = doc.add_paragraph()
        run = p.add_run(f"{full_label}")
        run.bold = True
        
        type_text = ""
        if 'select_one' in q_type: type_text = " [Single Choice]"
        elif 'select_multiple' in q_type: type_text = " [Multiple Choice]"
        elif 'text' in q_type: type_text = " [Text]"
        elif 'integer' in q_type or 'decimal' in q_type: type_text = " [Numeric]"
        elif 'calculate' in q_type: type_text = " [Calculation]"
        
        if type_text:
            p.add_run(type_text).italic = True

        details = []
        tech_cols = [
            ('required', 'Required'),
            ('relevance', 'Relevance'),
            ('calculation', 'Calculation'),
            ('constraint', 'Constraint'),
            ('constraint_message', 'Constraint Msg'),
            ('default', 'Default'),
            ('choice_filter', 'Choice Filter'),
            ('appearance', 'Appearance'),
            ('media', 'Media'),
            ('readonly', 'Read Only'),
            ('guidance_hint', 'Hint')
        ]

        for col_id, col_label in tech_cols:
            val = item.get(col_id)
            if val:
                details.append(f"{col_label}: {val}")

        ignored_keys = {'type', 'name', 'label::English', 'label::Hindi'}
        for key, val in item.items():
            if key not in ignored_keys and not key.startswith('Option_List::') and not any(k[0] == key for k in tech_cols):
                details.append(f"{key}: {val}")

        if details:
            detail_p = doc.add_paragraph(style='List Bullet')
            run_detail = detail_p.add_run("Technical Details: ")
            run_detail.bold = True
            run_tech = detail_p.add_run(" | ".join(details))
            run_tech.italic = True
            run_tech.font.size = Pt(9)

        if 'select_one' in q_type or 'select_multiple' in q_type:
            parts = q_type.split(' ')
            if len(parts) > 1:
                list_name = parts[1]
                options = choice_map.get(list_name, {})
                if options:
                    opt_items = list(options.items())[:25]
                    for code, details_opt in opt_items:
                        label = details_opt.get('Label', code) if isinstance(details_opt, dict) else details_opt
                        opt_p = doc.add_paragraph(style='List Bullet 2')
                        opt_p.add_run(f"{label} ({code})")
                    if len(opt_items) < len(options):
                        doc.add_paragraph(f"... and {len(options) - 25} more options omitted.", style='List Bullet 2')

    doc.save(output_path)
    return output_path
